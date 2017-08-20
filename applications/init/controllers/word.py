from gluon import contenttype
import docx
from cStringIO import StringIO
import datetime
from math import floor, ceil

DATE_3_MONTHS_BEFORE = (datetime.date.today() - datetime.timedelta(3 * 365 / 12))
DATE_6_MONTHS_BEFORE = (datetime.date.today() - datetime.timedelta(6 * 365 / 12))
DATE_9_MONTHS_BEFORE = (datetime.date.today() - datetime.timedelta(9 * 365 / 12))

# response.headers['Content-Type'] = contenttype.CONTENT_TYPE['.doc']
if not request.extension.lower() == 'doc':
    redirect(URL(f=request.function + ".doc", args=request.args, vars=request.get_vars))
response.headers['Content-Type'] = 'application/msword'


@auth.requires(URL.verify(request, hmac_key=MY_KEY, salt=session.MY_SALT, hash_vars=["type", "app_id"]))
# security to prevent SQL Injection attack
def tracking_chart():
    return dict(_type=request.get_vars["type"], **DOC_HEADER())


@auth.requires(URL.verify(request, hmac_key=MY_KEY, salt=session.MY_SALT, hash_vars=["app_id"]))
# security to prevent SQL Injection attack
def telephone_log():
    return dict(_type=request.get_vars["type"], **DOC_HEADER())


@auth.requires(URL.verify(request, hmac_key=MY_KEY, salt=session.MY_SALT, hash_vars=["app_id"]))
def discharge_poster():
    return DOC_HEADER()


@auth.requires(URL.verify(request, hmac_key=MY_KEY, salt=session.MY_SALT, hash_vars=["app_id"]))
def same_day_training_generic():
    return dict(EMR_NAME=APP.emr, **DOC_HEADER())


@auth.requires(URL.verify(request, hmac_key=MY_KEY, salt=session.MY_SALT, hash_vars=["app_id"]))
def er_ip_log():
    return dict(EMR_NAME=APP.emr, **DOC_HEADER())


def _docx_header():
    doc = docx.Document()
    doc_header = DOC_HEADER()
    p = doc.add_paragraph(doc_header["PRACTICE_NAME"])  # http://bit.ly/2rAnT4z
    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("{street}, {city}, {state} {zip}".format(
        street=doc_header["PRACTICE_STREET"],
        city=doc_header["PRACTICE_CITY"],
        state=doc_header["PRACTICE_STATE"],
        zip=doc_header["PRACTICE_ZIP"])
    )
    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Tel: {tel} | Fax: {fax}".format(
        tel=doc_header["PRACTICE_NUMBER"],
        fax=doc_header["PRACTICE_FAX"])
    )
    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    return doc


@auth.requires(URL.verify(request, hmac_key=MY_KEY, salt=session.MY_SALT, hash_vars=["app_id", "type"]))
def signin_sheet():
    output = StringIO()
    doc = _docx_header()
    h = doc.add_heading("Office Sign-in Sheet", 0)
    h.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Topic _________________________________            Date ____________", 3)
    doc.add_paragraph()
    p = doc.add_paragraph("Meeting [ %s ]       Training [ %s ]" % (
        "x" if request.get_vars["type"] == "meeting" else " ",
        "x" if request.get_vars["type"] == "training" else " "))
    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Position'
    hdr_cells[2].text = 'Signature'
    for item in range(10):
        row_cells = table.add_row().cells
        row_cells[0].text = "______________________________"
        row_cells[1].text = "____________________"
        row_cells[2].text = "_____________"
    # doc.add_paragraph()
    # doc.add_heading("Summary:", 3)
    # p = doc.add_paragraph("Duration: ______       Minutes [   ]       Hours [   ]")
    # p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    # doc.add_paragraph("_______________________________________________________________________________________________")
    # doc.add_paragraph("_______________________________________________________________________________________________")
    # doc.add_paragraph("_______________________________________________________________________________________________")
    # doc.add_paragraph("_______________________________________________________________________________________________")
    # doc.add_paragraph("_______________________________________________________________________________________________")
    doc.save(output)
    response.headers['Content-Type'] = 'application/msword'
    return output.getvalue()


@auth.requires(URL.verify(request, hmac_key=MY_KEY, salt=session.MY_SALT, hash_vars=["app_id"]))
def huddle_sheet():
    output = StringIO()
    doc = _docx_header()
    h = doc.add_heading("Daily Huddle Sheet", 0)
    h.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Date ____________", 3)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Checklist'
    hdr_cells[1].text = 'Huddle Agenda'
    agenda = [
        "Check for patients on the schedule that may require more time/assistance due to age, disability, personal "
        "demeanor, etc. who can help.",
        "Check for back-to-back lengthy appointments, such as physicals. How can they be worked around to prevent "
        "backlog?",
        "Are there openings which can be filled? Chronic now-shows? Any special instructions for the scheduler?",
        "Check over provider and staff schedule - does anyone need to leave early or break for a phone call or "
        "meeting?",
        "Lab results, test results, notes from other physicians, faxes -  are they signed and ready to be scanned "
        "into the patient's chart? Check for outstanding labs, referrals, and diagnostics.",
        "Review current schedule for high risk pts. High risk pts are highlighted in the EMR's OV screen."
    ]
    for item in agenda:
        row_cells = table.add_row().cells
        row_cells[0].text = "[    ]"
        row_cells[1].text = item

    doc.add_heading("Agenda:", 3)
    doc.add_paragraph("_______________________________________________________________________________________________")
    doc.add_paragraph("_______________________________________________________________________________________________")
    doc.add_paragraph("_______________________________________________________________________________________________")

    doc.save(output)
    return output.getvalue()


@auth.requires(URL.verify(request, hmac_key=MY_KEY, salt=session.MY_SALT, hash_vars=["app_id"]))
def policy_1a5_no_show():
    output = StringIO()
    doc = _docx_header()
    h = doc.add_heading("Monitor No-Show Rate Policy", 0)
    h.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    date = DATE_6_MONTHS_BEFORE
    doc.add_heading("%s/%s/%s" % (date.month, date.day, date.year), 3)
    doc.add_paragraph("")
    doc.add_paragraph("At the offices of {practice_name}, we have created the following workflow regarding "
                      "no show patients:".format(practice_name=APP.practice_name))

    doc.add_paragraph("The medical assistant will review the patients on the no show list and will call those patients "
                      "and make arrangements for a new appointment.  For those patients that are high priority will be "
                      "asked to come in the next day as a walk-in.",
                      style="ListNumber")

    doc.add_paragraph("In the cases where the patient is not available by telephone, the medical assistant will be "
                      "responsible for documenting the following in the patient's electronic health record (EHR):",
                      style="ListNumber")

    doc.add_paragraph("A 'missing you' letter will be sent to the patient address with a scanned copy in the "
                      "patient EHR.",
                      style="ListBullet2")
    doc.add_paragraph("If non-responsive, a follow-up letter, sent registered mail, return receipt will be sent "
                      "(5) business days later.",
                      style="ListBullet2")
    doc.save(output)
    return output.getvalue()


@auth.requires(URL.verify(request, hmac_key=MY_KEY, salt=session.MY_SALT, hash_vars=["app_id"]))
def report_1a5_no_show():
    output = StringIO()
    doc = _docx_header()
    h = doc.add_heading("Monitor No-Show Rate Policy", 0)
    h.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    date = datetime.date.today()
    doc.add_heading("%s/%s/%s" % (date.month, date.day, date.year), 3)
    doc.add_paragraph("")
    doc.add_paragraph("At the offices of {practice_name}, we have created the following workflow regarding "
                      "no show patients:".format(practice_name=APP.practice_name))

    doc.add_paragraph("The medical assistant will review the patients on the no show list and will call those patients "
                      "and make arrangements for a new appointment.  For those patients that are high priority will be "
                      "asked to come in the next day as a walk-in.")

    doc.save(output)
    return output.getvalue()


def policy_1a2_after_hours():
    output = StringIO()
    doc = _docx_header()
    h = doc.add_heading("PCMH 1A Factor 2: After-hours", 0)
    h.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    date = DATE_6_MONTHS_BEFORE
    doc.add_heading("%s/%s/%s" % (date.month, date.day, date.year), 3)
    doc.add_paragraph("")
    office_hours = db(db.office_hours.application == APP_ID).select()
    doc.add_paragraph("The office hours of {pc} is as follows:".format(pc=APP.practice_name))
    for hour in office_hours:
        doc.add_paragraph("{day_of_the_week} {start_time:%I}:{start_time:%M} {start_time:%p} - {end_time:%I}:{end_time:%M} {end_time:%p}".format(
            start_time=hour.start_time, end_time=hour.end_time, day_of_the_week=hour.day_of_the_week
        ))
    doc.save(output)
    return output.getvalue()

def report_1a1_same_day():
    output = StringIO()
    doc = _docx_header()
    h = doc.add_heading("PCMH 1A Factor 1: Same Day Appointments", 0)
    h.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    date = datetime.date.today()
    doc.add_heading("%s/%s/%s" % (date.month, date.day, date.year), 3)
    doc.add_paragraph("")
    doc.add_paragraph("At {practice_name}, we reserve time for same-day appointments, everyday when patients are seen. "
                      "The practice reserves time visibly on the scheduler on the following days"
                      .format(practice_name=APP.practice_name))
    doc.add_heading("[***ADD SCREENSHOT OF DENOMINATOR***]", 1)


@auth.requires(URL.verify(request, hmac_key=MY_KEY, salt=session.MY_SALT,
                          hash_vars=["app_id",
                                     "denominator", "hispanic", "non_hispanic", "black", "white", "native_american",
                                     "pacific_islander", "south_asian", "east_asian", "male", "female", "other",
                                     "english", "spanish", "chinese", "hindi", "bengali", "arabic", "african"]))
def report_2c_factor_1_2_demographics():
    output = StringIO()
    doc = _docx_header()
    h = doc.add_heading("Racial, Ethnic and Language Needs to Population", 0)
    h.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    date = datetime.date.today()
    doc.add_heading("%s/%s/%s" % (date.month, date.day, date.year), 3)
    doc.add_heading("[***ADD SCREENSHOT OF PATIENT SEARCH***]", 1)
    doc.add_heading("[***ADD SCREENSHOT OF DENOMINATOR***]", 1)
    doc.add_heading("Ethnicity", 1)
    ethnicities = ["hispanic", "non_hispanic"]
    ethnicity_table = doc.add_table(rows=len(ethnicities)+1, cols=4)
    ethnicity_hdr_cells = ethnicity_table.rows[0].cells
    ethnicity_hdr_cells[0].text = 'Subset'
    ethnicity_hdr_cells[1].text = 'Numerator'
    ethnicity_hdr_cells[2].text = 'Denominator'
    ethnicity_hdr_cells[3].text = 'Percent of Total Patients'
    for i, ethnicity in enumerate(ethnicities):
        ethnicity_cells = ethnicity_table.rows[1+i].cells
        ethnicity_cells[0].text = ethnicity.upper().replace("_", "-")
        ethnicity_cells[1].text = "%.0f" % (ceil(float(request.get_vars[ethnicity]) / 100.0 * float(request.get_vars["denominator"])))
        ethnicity_cells[2].text = request.get_vars["denominator"]
        ethnicity_cells[3].text = "%.0f%%" % (ceil(float(request.get_vars[ethnicity])))

    doc.add_heading("Race", 1)
    races = ["black", "white", "native_american", "pacific_islander", "south_asian", "east_asian"]
    race_table = doc.add_table(rows=len(races)+1, cols=4)
    race_hdr_cells = race_table.rows[0].cells
    race_hdr_cells[0].text = 'Subset'
    race_hdr_cells[1].text = 'Numerator'
    race_hdr_cells[2].text = 'Denominator'
    race_hdr_cells[3].text = 'Percent of Total Patients'
    for i, race in enumerate(races):
        print race, request.get_vars[race]
        race_cells = race_table.rows[1+i].cells
        race_cells[0].text = race.replace("_", " ").upper()
        race_cells[1].text = "%.0f" % (ceil(float(request.get_vars[race]) / 100.0 * float(request.get_vars["denominator"])))
        race_cells[2].text = request.get_vars["denominator"]
        race_cells[3].text = "%.0f%%" % (ceil(float(request.get_vars[race])))

    doc.add_heading("language", 1)
    languages = ["english", "spanish", "chinese", "hindi", "bengali", "arabic", "african"]
    language_table = doc.add_table(rows=len(languages)+1, cols=4)
    language_hdr_cells = language_table.rows[0].cells
    language_hdr_cells[0].text = 'Subset'
    language_hdr_cells[1].text = 'Numerator'
    language_hdr_cells[2].text = 'Denominator'
    language_hdr_cells[3].text = 'Percent of Total Patients'
    for i, language in enumerate(languages):
        print language, request.get_vars[language]
        language_cells = language_table.rows[1+i].cells
        language_cells[0].text = language.replace("_", " ").upper()
        language_cells[1].text = "%.0f" % (ceil(float(request.get_vars[language]) / 100.0 * float(request.get_vars["denominator"])))
        language_cells[2].text = request.get_vars["denominator"]
        language_cells[3].text = "%.0f%%" % (ceil(float(request.get_vars[language])))

    doc.add_heading("Gender", 1)
    genders = ["male", "female", "other"]
    gender_table = doc.add_table(rows=len(genders)+1, cols=4)
    gender_hdr_cells = gender_table.rows[0].cells
    gender_hdr_cells[0].text = 'Subset'
    gender_hdr_cells[1].text = 'Numerator'
    gender_hdr_cells[2].text = 'Denominator'
    gender_hdr_cells[3].text = 'Percent of Total Patients'
    for i, gender in enumerate(genders):
        gender_cells = gender_table.rows[1+i].cells
        gender_cells[0].text = gender.upper()
        gender_cells[1].text = "%.0f" % (ceil(float(request.get_vars[gender]) / 100.0 * float(request.get_vars["denominator"])))
        gender_cells[2].text = request.get_vars["denominator"]
        gender_cells[3].text = "%.0f%%" % (ceil(float(request.get_vars[gender])))

    doc.save(output)
    return output.getvalue()


def policy_2d9_staff_involvement():
    return