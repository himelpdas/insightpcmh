# -*- coding: utf-8 -*-
# this file is released under public domain and you can use without limitations

# -------------------------------------------------------------------------
# This is a sample controller
# - index is the default action of any application
# - user is required for authentication and authorization
# - download is for downloading files uploaded in the db (does streaming)
# -------------------------------------------------------------------------
from collections import OrderedDict
import json
import datetime
from gluon.tools import prettydate


db.logging.created_on.represent = db.application.created_on.represent = \
    lambda v, r: XML('<span title="%s">%s</span>' % (prettydate(getattr(r, 'created_on', None)),
                                                     getattr(r, 'created_on', None)))
db.application.modified_on.represent = lambda v, r: XML('<span title="%s">%s</span>' % (prettydate(
    getattr(r, 'modified_on', None)), getattr(r, 'modified_on', None)))

_role_to_permission = dict(app_managers="manage", contributors="contribute", admins="administrate", trainers="train",
                           masters=None)  # todo observers=observe


def test_email():
    response.view = os.path.join("templates", "email.html")
    return dict(summary="test", first_name="test", message="test", action_url="test", call_to_action="test")

def _disable_rbac_fields(*args):
    def _decorator(func):
        """when displaying a grid with something joined, sometimes the auth tables joined are not necessary"""
        def inner():  # http://stackoverflow.com/questions/19673284/how-do-i-get-list-of-field-objects-in-a-table-in-web2py
            for rbac in args:
                for field in rbac:
                    field.readable = False
                    field.writable = False
            return func
        return inner()
    return _decorator


def index():
    # web2py performs inner joins automatically and transparently when the query links two or more tables
    response.title = "Insight PCMH"
    response.subtitle = " - Patient-Centered Medical Home Transformation for Practices"
    return dict()


@_disable_rbac_fields(db.auth_membership, db.auth_permission, db.auth_group)
@auth.requires_login()
def dash():
    # web2py performs inner joins automatically and transparently when the query links two or more tables
    response.title = "PCMH Dashboard"
    return dict()

@auth.requires(URL.verify(request, hmac_key=MY_KEY, salt=session.MY_SALT,
                          hash_vars=["revoke_participant", "permission", "row_id"]))
def revoke_user():
    e = request.get_vars["revoke_participant"]
    p = request.get_vars["permission"]
    r = request.get_vars["row_id"]
    del request.get_vars["row_id"]
    del request.get_vars["revoke_participant"]
    del request.get_vars["permission"]
    del request.get_vars["_signature"]

    logger.warn("Revoking, user id, permission, row, self-group:\n%s, %s, %s, %s" % (e, p, r, auth.id_group("user_%s" % e)))

    auth.del_permission(auth.id_group("user_%s" % e), p, "application", r)  # 0 means user_1

    session.flash = "Revoked user from application ID%s" % r
    redirect(URL("dash.html", args=request.args, vars=request.get_vars))


@auth.requires(URL.verify(request, hmac_key=MY_KEY, salt=session.MY_SALT,
                      hash_vars=["assign_participant", "permission", "row_id"]))
def assign_user():

    e = request.get_vars["assign_participant"]
    p = request.get_vars["permission"]
    r = request.get_vars["row_id"]
    del request.get_vars["row_id"]
    del request.get_vars["assign_participant"]
    del request.get_vars["permission"]
    del request.get_vars["_signature"]

    auth.add_permission(auth.id_group("user_%s" % e), p, "application", r)  # 0 means user_1
    #if not auth.has_membership(role="contributors", user_id=e):
    #    auth.add_membership(role="contributors", user_id=e)

    session.flash = "Assigned user to application ID%s" % r
    redirect(URL("dash.html", args=request.args, vars=request.get_vars))


def _assigned_column(row):
    row.id = getattr(row, 'id', None) or getattr(row.application, 'id', None)  # can be from join or regular query

    assert row.id, "expected row.id"

    participators = PARTICIPATORS(row)
    #
    # print participators
    participator_widgets = []
    for participator in participators:
        base = dict(
            c_id=participator.auth_user.id,
            c_fn=participator.auth_user.first_name.capitalize(),
            c_ln=participator.auth_user.last_name.capitalize(),
            c_email=participator.auth_user.email,
        )
        base.update(
            c_posessive="%s's" % base["c_fn"],
            c_name="%s %s (%s)" % (base['c_fn'], base['c_ln'], base['c_id']),
            c_name_html="<a href='mailto:%s'>%s %s</a> <span class='text-muted'>(%s)</span>" %
                        (base['c_email'], base['c_fn'], base['c_ln'], base['c_id']),
            c_acronym="%s%s%s" % (participator.auth_user.first_name.capitalize()[0],
                                  participator.auth_user.last_name.capitalize()[0],
                                  participator.auth_user.id),
        )
        if auth.has_membership(user_id=participator.auth_user.id, role="trainers") and \
                        participator.auth_permission.name == "train":
            participator_widgets.append(dict(color="danger", title="trainer", permission="train",
                                            c_title=base['c_name']+" (Trainer)",
                                            c_title_html=base['c_name_html'] +
                                            " <span class='text-danger'>(Trainer)</span>", **base))
        elif auth.has_membership(user_id=participator.auth_user.id, role="app_managers") and \
                        participator.auth_permission.name == "manage":
            participator_widgets.append(dict(color="warning", title="app_manager", permission="manage",
                                            c_title=base['c_name']+" (App Manager)",
                                            c_title_html=base['c_name_html'] +
                                            " <span class='text-warning'>(App Manager)</span>", **base))  # it is possible to be two roles
        elif participator.auth_permission.name == "contribute":
            is_contrib = auth.has_membership(user_id=participator.auth_user.id, role="contributors")
            fake_title = "Contributor" if is_contrib else "Visitor"
            fake_color = "success" if is_contrib else "primary"
            participator_widgets.append(dict(color=fake_color, title="contributor",
                                             permission="contribute", c_title=base['c_name'] + " (%s)" % fake_title,
                                             c_title_html=base['c_name_html'] +
                                             " <span class='text-%s'>(%s)</span>" % (fake_color, fake_title), **base))

    for participator_widget in participator_widgets:
        revoke_id = "revoke_participant_%s_%s_%s" % (participator_widget["title"], participator_widget["c_id"], row.id)
        widget_script = """
        {select}<script>
            $(document).ready(function(){{
            $("#{revoke_id}").val('').multiselect({{
                nonSelectedText: '{c_acronym}',
                onChange: function(option, checked, select) {{
                    window.location = '{url}';
                }},
                buttonClass: 'btn btn-sm btn-{color}', disableIfEmpty: true, enableHTML: true}});
            }})
        </script>"""  # option
        widget_a = XML("<button class='btn btn-sm btn-{color}' "
                                "title=\"{c_name_html}\">{c_acronym}</button>".format(
                color=participator_widget["color"],
                c_acronym=participator_widget["c_acronym"],
                c_name_html=participator_widget["c_title_html"], # todo change to popover
        ))
        widget_b = XML(widget_script.format(
            c_acronym=participator_widget["c_acronym"],
            revoke_id=revoke_id,
            color=participator_widget["color"],
            url=URL("revoke_user.html",
                    vars=dict(
                        revoke_participant=participator_widget["c_id"],
                        permission=participator_widget["permission"],
                        row_id=row.id,
                        **request.get_vars),
                    hmac_key=MY_KEY, salt=session.MY_SALT,
                    hash_vars=["revoke_participant", "permission", "row_id"]
                    ),
            select=SELECT(OPTGROUP(OPTION("Revoke %s access to this application" % participator_widget['c_posessive']),
                                   _label=participator_widget['c_title_html'] + ":"),
                          _id=revoke_id,
                          _class="assigned", )
        ))

        if IS_ADMIN or IS_MASTER:
            widget = widget_b  # only admins can revoke
        else:
            widget = widget_a

        #assert(widget, "expected a widget!")

        participator_widget.update(dict(
            widget=widget
        ))

    colleagues = db((db.auth_group.id == db.auth_membership.group_id) &
                    (db.auth_group.role.belongs("trainers", "app_managers", "contributors")) &
                    (db.auth_user.id == db.auth_membership.user_id)).select(orderby=db.auth_user.first_name|db.auth_group.role)

    employee_options = []
    contributor_options = []
    assign_urls = {}

    title = e_title = e_title_html = permission = None
    for colleague in colleagues:
        e_id = colleague.auth_user.id
        e_fn = colleague.auth_user.first_name.capitalize()
        e_ln = colleague.auth_user.last_name.capitalize()
        e_email = colleague.auth_user.email
        e_name = "%s %s (%s)" % (e_fn, e_ln, e_id)
        e_name_html = "%s %s <span class='text-muted'>(%s)</span>" % (e_fn, e_ln, e_id)

        if colleague.auth_group.role == "trainers":
            title = "trainer"
            permission = "train"
            e_title = e_name + " (trainer)"
            e_title_html = e_name_html + " <span class='text-danger'>(Trainer)</span>"        
        
        if colleague.auth_group.role == "app_managers":
            title = "app_manager"
            permission = "manage"
            e_title = e_name + " (App Mananger)"
            e_title_html = e_name_html + " <span class='text-warning'>(App Manager)</span>"
            
        if colleague.auth_group.role == "contributors":
            title = "contributor"
            permission = "contribute"
            e_title = e_name + " (Contributor)"
            e_title_html = e_name_html + " <span class='text-success'>(Contributor)</span>"

        assign_id = "%s_%s" % (title, e_id)
        assign_urls[assign_id] = URL("assign_user.html",
                                     vars=dict(
                                         assign_participant=e_id,
                                         permission=permission,
                                         row_id=row.id,
                                         **request.get_vars),
                                     hmac_key=MY_KEY, salt=session.MY_SALT,
                                     hash_vars=["assign_participant", "permission", "row_id"]
                                     )
        # if not already in revoke buttons
        if not "%s_%s" % (title, e_id) in map(lambda w: "%s_%s" % (w["title"], w["c_id"]), participator_widgets):
            _bucket = employee_options
            if colleague.auth_group.role == "contributors":
                _bucket = contributor_options
            _bucket.append(OPTION(e_title_html, _value=assign_id))

    colleagues_and_participants = set(map(lambda e: e.auth_user.id, colleagues) + map(lambda e: e.auth_user.id, 
                                                                                    participators))
    # a_week_ago = request.now - datetime.timedelta(days=7)
    recent = db(db.auth_user.created_on > 0).select(orderby=~db.auth_user.modified_on, limitby=(0, 50))
    recent.exclude(lambda r: r.id in colleagues_and_participants or r.is_insight)
    recent_options = []
    
    for new in recent:
        n_id = new.id
        n_fn = new.first_name.capitalize()
        n_ln = new.last_name.capitalize()
        n_email = new.email
        n_name = "%s %s (%s)" % (n_fn, n_ln, n_id)
        n_name_html = "%s %s <span class='text-muted'>(%s) (%s)</span>" % (n_fn, n_ln, n_id, n_email)
        assign_id = "new_%s" % n_id
        assign_urls[assign_id] = URL("assign_user.html",
                                     vars=dict(
                                         assign_participant=n_id,
                                         permission="contribute",
                                         row_id=row.id,
                                         **request.get_vars),
                                     hmac_key=MY_KEY, salt=session.MY_SALT,
                                     hash_vars=["assign_participant", "permission", "row_id"]
                                     )
        recent_options.append(OPTION(n_name_html, _value=assign_id))

    employee_optgroup = OPTGROUP(*employee_options, _label="Assign an employee:")
    contributor_optgroup = OPTGROUP(*contributor_options, _label="Assign a contributor:")
    recent_optgroup = OPTGROUP(*recent_options, _label="Assign a recent registrant:")

    assign_optgroups = []
    enable_participant_assign_select = False
    if IS_ADMIN or IS_MASTER:
        if recent_options:
            assign_optgroups.append(recent_optgroup)
        if contributor_options:  # no need to show opt-group if there are no options available
            assign_optgroups.append(contributor_optgroup)
        if employee_options:
            assign_optgroups.append(employee_optgroup)
        enable_participant_assign_select = True

    all_widgets = map(lambda e: e['widget'], participator_widgets)

    if enable_participant_assign_select:
        all_widgets.append(XML("""{select}<script>
                    $(document).ready(function(){{
                    urls_{row_id} = {urls};
                    $("#assign_participant_{row_id}").val('').multiselect({{
                        nonSelectedText: '<span class="glyphicon glyphicon-plus"></span>',
                        maxHeight: 200,
                        enableFiltering: true,
                        enableCaseInsensitiveFiltering: true,
                        onChange: function(option, checked, select) {{
                            window.location = urls_{row_id}[$(option).val()];
                        }},
                        buttonClass: 'btn btn-sm btn-info', disableIfEmpty: true, enableHTML: true}});
                    }})
                </script>""".format(  # option #http://bit.ly/2g9PyCl select remove default choice, then style with multiselect
            select=SELECT(
                *reversed(assign_optgroups),
                _id="assign_participant_%s" % row.id,
                _class="assign_participants"
            ),
            row_id=row.id,
            urls=json.dumps(assign_urls)
        )))  # , _multiple="multiple")  # only one choice

    container = DIV(*all_widgets, _style="display:flex")
    return container


#@_disable_rbac_fields
# @auth.requires_signature()
def load_apps_grid():
    # db.application.modified_by.readable = True
    # db.application.created_by.readable = True
    # db.application.created_on.readable = True
    # db.application.modified_on.readable = True

    db.application.modified_by.readable = True
    db.application.modified_on.readable = True
    db.application.created_on.readable = True
    db.application.owner_id.readable = True

    links = [dict(header='',  # header is col title
                  body=lambda row:
                  A(SPAN(_class="glyphicon glyphicon-play"),
                    _class="btn btn-sm btn-default",
                    _title="Start Portal",
                    _href=URL("2014", 'index.html', args=[0],  # todo- set 2014/2017 standards here table may or may not be joined
                              vars=dict(app_id=getattr(row, "application", row).id))))]

    if not IS_STAFF:  # add contributor
        db.application.status.writable = False
    if IS_MASTER or IS_ADMIN:
        db.application.owner_id.writable = True
    if IS_MASTER:  # remove not after testing non-master mode
        my_apps_grid = db(db.application.id > 0)
    else:
        my_group_id = auth.id_group("user_%s" % auth.user.id)
        my_apps_disinct = db((db.application.id == db.auth_permission.record_id) &  # same application id will show up
                             # twice because multiple permissions of same user can be set for the same application (i.e.
                             # when you see HD1 XXX HD1 in master mode *WARNING*
                     (db.auth_permission.name.belongs(["manage", "contribute", "administrate", "train"])) &
                     (db.auth_permission.group_id == my_group_id)).select(groupby=db.application.id)  # distinct gives near "ON" operational error, just use groupby http://bit.ly/2h0Ou3Z
        # groupby http://bit.ly/2h0Ou3Z

        logger.info(my_apps_disinct)

        my_apps_grid = db.application.id.belongs(map(lambda r: r.application.id, my_apps_disinct))  # will have to
        # double query because grid does not have distinct and groupby disables CUD

    links.append(dict(
        header="Participants",  # can use SPAN
        body=_assigned_column
    ))

    app_grid = SQLFORM.grid(my_apps_grid,
                            onvalidation=_app_onvalidation,
                            oncreate=_app_oncreate,
                            create=IS_ADMIN or IS_MASTER or IS_CONTRIB,
                            formname="load_apps_grid",
                            links=links,
                            deletable=IS_ADMIN or IS_MASTER or IS_CONTRIB,
                            editable=IS_ADMIN or IS_MASTER or IS_CONTRIB,
                            # groupby=db.application.id, # groupby by itself behaves like distinct http://bit.ly/2h0Ou3Z
                            field_id=db.application.id,
                            links_placement='left')

    return dict(app_grid=app_grid)


@auth.requires(URL.verify(request, hmac_key=MY_KEY, salt=session.MY_SALT, hash_vars=["group", "action",
                                                                                           "user_id"]))
def add_remove_user():
    u = request.get_vars["user_id"]
    a = request.get_vars["action"]
    g = request.get_vars["group"]
    del request.get_vars["group"]
    del request.get_vars["action"]
    del request.get_vars["user_id"]
    del request.get_vars["_signature"]

    user = db(db.auth_user.id == u).select().last()

    assert user, "expected user! (user deleted?)"
    assert g in _without_keys(_role_to_permission, "masters"), "Invalid role (web parameter tampering?)" # not needed because of auth.signature

    if a == "+":
        auth.add_membership(role=g, user_id=u)
        word = "added"
        direction = "to"
    else:
        print g
        print auth.id_group("user_%s" % u)
        auth.del_membership(role=g, user_id=u)
        db((db.auth_permission.group_id == auth.id_group("user_%s" % u)) &
            (db.auth_permission.name == _role_to_permission[g])).delete()
        word = "removed"
        direction = "from"

    session.flash = '%s was %s %s the group "%s"' % (user.first_name.capitalize(), word, direction, g.replace("_"," "))

    if user.email in MASTER_EMAILS:
        session.flash = "Don't even think about it! ;)"

    redirect(URL("dash.html", args=request.args, vars=request.get_vars))


def _employee_group_links(row):
    _p = SPAN(_class="glyphicon glyphicon-plus")  # plus or minus
    _m = SPAN(_class="glyphicon glyphicon-minus")

    if row.is_insight:
        groups = [("admins", "info"), ("app_managers", "warning"), ("trainers", "danger")]
    else:
        groups = [("contributors", "success")]
    all_links = []
    for each in groups:
        group = each[0]
        group_name = group.capitalize().replace("_", " ")
        color = each[1]
        action = "-" if auth.has_membership(role=group, user_id=row.id) else "+"
        sign = _m if action == "-" else _p
        btn_color = "btn-secondary"
        confirm_msg = ""
        if action != "+":
            confirm_msg = " All assignments for this user will be lost!"
            btn_color = "btn-%s" % color
        label = XML("%s %s" % (sign,
                               group_name[:-1] if not group == "contributors" else "Contributor (Write Access)")
                    )  # take out the s at the end
        all_links.append(A(label, _onclick="if(!confirm('Are you sure?%s')){event.preventDefault()}" % confirm_msg,  # http://bit.ly/2hM7Cbk
                           _class="btn btn-sm %s" % btn_color,
                           _href=URL("add_remove_user.html",
                                     vars=dict(
                                         user_id=row.id,
                                         group=group,
                                         action=action,
                                         **request.get_vars),
                                     hmac_key=MY_KEY, salt=session.MY_SALT,
                                     hash_vars=["user_id", "group", "action"]
                                     ),
                           _title=(("Remove %s from " if action == "-" else "Add %s to ")
                                   % row.first_name.capitalize()) + group_name
                           )
                         )
    return DIV(*all_links, _style="display:flex")


def load_logs_grid():
    db.logging.owner_id.default=auth.user.id
    db.logging.owner_id.writable=False
    db.logging.created_on.readable=True
    db.logging.created_by.readable=True
    db.application.id.represent = lambda v, r: "%s (%s)" % (r.application.practice_name, r.application.id)

    if IS_MASTER or IS_ADMIN:
        db.logging.owner_id.readable = True
        query = db(db.logging.id > 0)
    else:
        db.logging.owner_id.readable = False
        query = db((db.logging.owner_id == db.auth_user.id) & (db.auth_user.id == auth.user.id))

    logs_grid = SQLFORM.grid(
        query,
        formname="load_logs_grid",
        deletable=False if not IS_MASTER else True,
        editable=False if not IS_MASTER else True,
        fields=[db.logging.id, db.logging.application, db.logging.difficulty, db.logging.description,
                db.logging.created_by, db.logging.created_on],
        #links=links,
        #onupdate=_user_onupdate,
        #oncreate=_user_oncreate,
        orderby=~db.logging.id,  # show latest first
        field_id=db.logging.id,
        # user_signature=False,  # this is handled by the controller
        links_placement='left')


    return dict(logs_grid=logs_grid)

#@_disable_rbac_fields
# @auth.requires_signature()
def load_users_grid():
    links = []

    links.append(dict(
        header="Set roles",  # can use SPAN
        body=_employee_group_links
    ))

    users_grid = SQLFORM.grid(db.auth_user,
                              formname="load_users_grid",
                              links=links,
                              onupdate=_user_onupdate,
                              oncreate=_user_oncreate,
                              orderby=~db.auth_user.is_insight | db.auth_user.id,  # show insight first
                              field_id=db.auth_user.id,
                              #user_signature=False,  # this is handled by the controller
                              links_placement='left')

    return dict(users_grid=users_grid)


def user():
    """
    exposes:
    http://..../[app]/default/user/login
    http://..../[app]/default/user/logout
    http://..../[app]/default/user/register
    http://..../[app]/default/user/profile
    http://..../[app]/default/user/retrieve_password
    http://..../[app]/default/user/change_password
    http://..../[app]/default/user/bulk_register
    use @auth.requires_login()
        @auth.requires_membership('group name')
        @auth.requires_permission('read','table name',record_id)
    to decorate functions that need access control
    also notice there is http://..../[app]/appadmin/manage/auth to allow administrator to manage users
    """
    response.title = (request.args(0) or request.function).capitalize().replace("_", " ")
    return dict(form=auth())


#auth.requires(not auth.has_membership(user_id=getattr(auth.user, "id", None), role="trainers") and
#               not auth.has_membership(user_id=getattr(auth.user, "id", None), role="admins")
#               , requires_login=True)
def _app_onvalidation(form):
    if not IS_STAFF:
        form.vars.owner_id = auth.user.id
    if form.vars.application_size == "Corporate":
        corporate_apps = db(db.application.authorized_representative == form.vars.authorized_representative).select()
        if form.vars.largest_practice:
            for app in corporate_apps:
                if app.largest_practice:
                    form.errors.largest_practice = "Largest practice under the authorized representative already exists!"


def _user_onupdate(form):  # todo revoke all permissions
    id = form.vars.id
    self_group = auth.id_group("user_%s" % id)
    # auth.del_permission(auth.id_group("user_%s" % e), p, "application", r)  # 0 means user_1
    if not form.vars.is_insight:
        groups = _without_keys(_role_to_permission, "contributors")
        for role in groups:
            permission = groups[role]
            auth.del_membership(role=role, user_id=id)
            if permission:
                db((db.auth_permission.group_id == self_group) &
                    (db.auth_permission.name == permission)).delete()
    else:
        auth.del_membership(role="contributors", user_id=id)
        db((db.auth_permission.group_id == self_group) &
            (db.auth_permission.name == "contribute")).delete()


def _user_oncreate(form):
    id = form.vars.id
    self_group = "user_%s" % id
    auth.add_group(self_group, description="Group for user %s. Created in admin panel" % self_group)
    auth.add_membership(role=self_group, user_id=id)


def _app_oncreate(form):
    app_id = form.vars.id

    if not IS_STAFF:
        auth.add_permission(0, "contribute", 'application', app_id)  # 0 means user_1
        #auth.add_membership(role="contributor", user_id=auth.user_id)

    admins = db((db.auth_group.id == db.auth_membership.group_id) &
                (db.auth_group.role == "admins") &
                (db.auth_user.id == db.auth_membership.user_id)
                ).select()

    for admin in admins:
        auth.add_permission(auth.id_group("user_%s" % admin.auth_user.id), "administrate", 'application', app_id)  # 0 means user_1
    #auth.add_permission(auth.id_group("masters"), "administrate", 'application', app_id)
    # app_url = URL('application', vars=dict(app_id=form.vars.id), hmac_key=MY_KEY)
    #redirect(URL(0, "index.html", vars=dict(app_id=form.vars.id)))  # redir user to his app


@cache.action()
def download():
    """
    allows downloading of uploaded files
    http://..../[app]/default/download/[filename]
    """
    return response.download(request, db)


def call():
    """
    exposes services. for example:
    http://..../[app]/default/call/jsonrpc
    decorate with @services.jsonrpc the functions to expose
    supports xml, json, xmlrpc, jsonrpc, amfrpc, rss, csv
    """
    return service()

import docx
from cStringIO import StringIO
def test_word():
    output = StringIO()
    doc = docx.Document()
    doc.add_heading('Document Title', 0)

    p = doc.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    doc.add_heading('Heading, level 1', level=1)
    doc.add_paragraph('Intense quote', style='IntenseQuote')

    doc.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    doc.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )
    doc.save(output)
    response.headers['Content-Type'] = 'application/msword'
    return output.getvalue()
